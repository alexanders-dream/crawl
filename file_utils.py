# file_utils.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def convert_to_docx(content: str) -> bytes:
    """Convert markdown content to DOCX format"""
    document = Document()
    content = content.strip()
    
    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    
    for paragraph in paragraphs:
        if paragraph.startswith('# '):
            # Heading
            heading = document.add_heading(paragraph[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.startswith('## '):
            # Subheading
            heading = document.add_heading(paragraph[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.startswith('- '):
            # Bullet point
            document.add_paragraph().add_run(paragraph).font.size = Pt(10)
        else:
            # Regular paragraph
            document.add_paragraph(paragraph)
            
    # Save to bytes
    file_bytes = BytesIO()
    document.save(file_bytes)
    file_bytes.seek(0)
    return file_bytes.read()